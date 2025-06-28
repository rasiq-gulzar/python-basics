from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import os
from docx.oxml.ns import qn

# Create a new Word document
doc = Document()

# Set up title page
doc.add_heading('القرآن الكريم', level=1).alignment = 1
doc.add_paragraph("Al-Quran Al-Kareem").alignment = 1
doc.add_paragraph("\n\nPrepared by Rasiq Gulzar").alignment = 1
doc.add_page_break()

# Add Surah Title Style
style = doc.styles.add_style('SurahTitle', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = 'Traditional Arabic'
style.font.size = Pt(18)
style.font.bold = True

# Add Arabic Verse Style
style = doc.styles.add_style('ArabicText', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = 'Traditional Arabic'
style.font.size = Pt(16)
style.font.bold = False
style.font.color.rgb = RGBColor(0, 0, 139)

# Add Translation Style
style = doc.styles.add_style('TranslationText', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.font.italic = True

# Example Content
doc.add_paragraph("سورة الفاتحة", style='SurahTitle').alignment = 1

# Add verses with translation in table
table = doc.add_table(rows=1, cols=2)
table.autofit = True

# First verse
arabic_cell = table.rows[0].cells[0]
arabic_para = arabic_cell.paragraphs[0]
arabic_para.text = "الْحَمْدُ لِلَّهِ رَبِّ الْعَالَمِينَ"
arabic_para.style = 'ArabicText'

# Fix: Set RTL direction properly
tcPr = arabic_cell._element.get_or_add_tcPr()
textDirection = OxmlElement('w:dir')
textDirection.set(qn('w:val'), 'rtl')
tcPr.append(textDirection)

trans_cell = table.rows[0].cells[1]
trans_para = trans_cell.paragraphs[0]
trans_para.text = "Praise be to Allah, Lord of all the worlds"
trans_para.style = 'TranslationText'

# Second verse
row_cells = table.add_row().cells
arabic_para = row_cells[0].paragraphs[0]
arabic_para.text = "الرَّحْمَٰنِ الرَّحِيمِ"
arabic_para.style = 'ArabicText'

# Fix: Set RTL direction for second verse
tcPr = row_cells[0]._element.get_or_add_tcPr()
textDirection = OxmlElement('w:dir')
textDirection.set(qn('w:val'), 'rtl')
tcPr.append(textDirection)

trans_para = row_cells[1].paragraphs[0]
trans_para.text = "The Most Compassionate, the Most Merciful"
trans_para.style = 'TranslationText'

# Save document
try:
    documents_path = os.path.join(os.path.expanduser('~'), 'Documents')
    output_path = os.path.join(documents_path, 'Quran_Template.docx')
    os.makedirs(documents_path, exist_ok=True)
    doc.save(output_path)
    print(f"Document saved successfully at: {output_path}")
except Exception as e:
    print(f"Error saving document: {str(e)}")